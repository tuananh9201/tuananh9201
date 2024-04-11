from docx import Document
from docx.shared import Inches, Pt
from os import walk
import os
import chardet


ALLOW_FILES_EXTENSIONS = ['.py', '.tsx','.ts', '.scss', '.sass','.css']

IGNORE_FOLDERS = ["node_modules", ".git","__pycache__"]

IGNORE_FIDLES = ["main.py","main2.py"]

IS_DEV = False
STOP = False
"""
Recursively walks through a directory and its subdirectories, and returns a list of all files with allowed extensions, excluding files and folders that should be ignored.

Args:
    directory (str): The directory to start the search from.
    allowed_extensions (list): A list of file extensions to include in the search.
    ignore_folders (list): A list of folder names to exclude from the search.
    ignore_files (list): A list of file names to exclude from the search.

Returns:
    list: A list of file paths that match the search criteria.
"""


def write_to_docx(path, docx_filename):
  """
  Writes file information to a docx file.

  Args:
      path: The directory path to process.
      docx_filename: The name of the docx file to create.
  """
  document = Document()
  count_file = 0

  # Loop through files in the directory
  for root, _, files in walk(path):
    for filename in files:
      #Skip files in ignored folders or wrong extension
        # Skip files in ignored folders or wrong extension or in ignore_files list
        if any(ignore_folder in root for ignore_folder in IGNORE_FOLDERS) or \
            not filename.endswith(tuple(ALLOW_FILES_EXTENSIONS)) or \
            filename in IGNORE_FIDLES:
            continue
        try:
           # Get relative path
            relative_path = os.path.relpath(os.path.join(root, filename), path)
            
            # Add heading and content
            paragraph = document.add_paragraph()
            heading = paragraph.add_run(relative_path)
            heading.bold = True
            paragraph.add_run("\n")
            document.add_paragraph()


            with open(os.path.join(root, filename), 'rb') as f:
                rawdata = f.read()
                document.add_paragraph()
                result = chardet.detect(rawdata)
                encoding  = result['encoding']
                with open(os.path.join(root, filename), 'r', encoding=encoding) as f:
                    content = f.read()
                    paragraph.add_run(content)
                    document.add_paragraph()
        except Exception as ex:
            print('----------------- Exception ---------------------')
            print(ex)
            print('----------------- End Exception ---------------------')
            continue
        finally:
            count_file += 1
        if IS_DEV and count_file>=5:
            print('------STOP------')
            document.save(docx_filename)
            print(f"STOP --- File information written to {docx_filename}")
            return
  # Save the document
  document.save(docx_filename)
  print(f"File information written to {docx_filename}")

# Get current directory (replace with your desired path)
current_dir = os.getcwd()

# Specify docx filename
docx_filename = "file_info.docx"

# Write to docx
write_to_docx(current_dir, docx_filename)
